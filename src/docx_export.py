"""Export AI reports as Microsoft Word documents."""

from docx import Document


def export_docx(report_text: str, output_path: str):
    doc = Document()
    doc.add_heading('Weekly AI Research Report', level=1)
    for line in report_text.splitlines():
        if line.strip():
            doc.add_paragraph(line)
    doc.save(output_path)
    return output_path
